"""
模块9：多格式输出导出器
- HTML → PDF 转换
- Markdown → Word (.docx) 转换
- 增量修改支持（版本管理+变更日志）
- 数据隐私声明生成
"""

import json
import shutil
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
from dataclasses import dataclass

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import OUTPUT_DIR, DATA_RETENTION_DAYS
from utils.helpers import (
    ensure_dir, write_text_file, save_json, load_json,
    format_change_log, sanitize_filename
)
from modules.content_generator import GeneratedDocument
from modules.layout_engine import LayoutEngine


@dataclass
class ExportResult:
    """导出结果"""
    project_name: str
    output_dir: Path
    markdown_path: Optional[Path] = None
    html_path: Optional[Path] = None
    pdf_path: Optional[Path] = None
    docx_path: Optional[Path] = None
    missing_report_path: Optional[Path] = None
    quality_report_path: Optional[Path] = None
    privacy_statement_path: Optional[Path] = None
    version: int = 1
    exported_at: str = ""


class OutputExporter:
    """
    多格式输出导出器

    职责：
    1. Markdown中间格式保存
    2. HTML美化版生成
    3. PDF提交版生成
    4. Word可编辑版生成
    5. 版本管理与变更日志
    6. 数据隐私声明
    """

    def __init__(self, output_root: Path = None):
        self.output_root = Path(output_root) if output_root else OUTPUT_DIR

    def export_all(
        self,
        document: GeneratedDocument,
        layout_engine: LayoutEngine,
        quality_report_text: str = "",
        version: int = 1,
        change_log: Optional[List[Dict[str, str]]] = None,
    ) -> ExportResult:
        """
        导出所有格式

        Args:
            document: 生成的策划书
            layout_engine: 排版引擎（含视觉样式）
            quality_report_text: 质量检查报告文本
            version: 版本号
            change_log: 变更日志（增量修改时使用）

        Returns:
            ExportResult: 所有导出文件的路径
        """
        # 创建项目专属输出目录
        project_dir_name = sanitize_filename(
            f"{document.project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_v{version}"
        )
        prj_output_dir = self.output_root / project_dir_name
        ensure_dir(prj_output_dir)

        result = ExportResult(
            project_name=document.project_name,
            output_dir=prj_output_dir,
            version=version,
            exported_at=datetime.now().isoformat(),
        )

        full_text = document.get_full_text()

        # 1. Markdown 中间格式
        md_path = prj_output_dir / "final_plan.md"
        write_text_file(md_path, full_text)
        result.markdown_path = md_path

        # 2. HTML 美化版
        html_content = layout_engine.render_html(
            markdown_content=full_text,
            project_name=document.project_name,
            competition_name=document.competition_name,
        )
        html_path = prj_output_dir / "final_plan.html"
        write_text_file(html_path, html_content)
        result.html_path = html_path

        # 3. PDF 提交版（尝试使用weasyprint）
        try:
            result.pdf_path = self._export_pdf(html_content, prj_output_dir)
        except Exception as e:
            print(f"[OutputExporter] PDF导出失败（可能未安装weasyprint）: {e}")
            print(f"[OutputExporter] 请手动从HTML转换PDF，或运行: pip install weasyprint")

        # 4. Word 可编辑版
        try:
            result.docx_path = self._export_docx(full_text, document.project_name, prj_output_dir)
        except Exception as e:
            print(f"[OutputExporter] Word导出失败: {e}")

        # 5. 缺失项清单
        missing_report = document.get_missing_report()
        missing_path = prj_output_dir / "missing_checklist.md"
        write_text_file(missing_path, missing_report)
        result.missing_report_path = missing_path

        # 6. 质量检查报告
        if quality_report_text:
            qr_path = prj_output_dir / "quality_report.md"
            write_text_file(qr_path, quality_report_text)
            result.quality_report_path = qr_path

        # 7. 变更日志
        if change_log:
            cl_path = prj_output_dir / "change_log.md"
            write_text_file(cl_path, format_change_log(change_log))
        else:
            cl_path = prj_output_dir / "change_log.md"
            write_text_file(cl_path, format_change_log([]))

        # 8. 数据隐私声明
        privacy_path = prj_output_dir / "DATA_PRIVACY.txt"
        write_text_file(privacy_path, self._generate_privacy_statement(document.project_name))
        result.privacy_statement_path = privacy_path

        # 9. 客户补充资料引导问卷（覆盖所有缺失类型，不仅仅是财务）
        if document.missing_sections:
            q_path = prj_output_dir / "client_supplement_guide.md"
            write_text_file(q_path, self._generate_comprehensive_questionnaire(
                document.project_name, document.missing_sections))
            # 同样保留财务专用问卷
            financial_missing = [m for m in document.missing_sections
                               if any(kw in m for kw in ["财务", "成本", "营收", "融资", "估值"])]
            if financial_missing:
                fq_path = prj_output_dir / "financial_questionnaire.md"
                write_text_file(fq_path, self._generate_financial_questionnaire(document.project_name))

        # 9. 保存元数据
        metadata = {
            "project_name": document.project_name,
            "competition_name": document.competition_name,
            "template_used": document.template_used,
            "total_word_count": document.total_word_count,
            "version": version,
            "exported_at": result.exported_at,
            "formats_exported": {
                "markdown": str(result.markdown_path) if result.markdown_path else None,
                "html": str(result.html_path) if result.html_path else None,
                "pdf": str(result.pdf_path) if result.pdf_path else None,
                "docx": str(result.docx_path) if result.docx_path else None,
            }
        }
        save_json(prj_output_dir / "metadata.json", metadata)

        return result

    def incremental_update(
        self,
        previous_result: ExportResult,
        updated_document: GeneratedDocument,
        layout_engine: LayoutEngine,
        changes: List[Dict[str, str]],
    ) -> ExportResult:
        """
        增量修改：只更新有变化的部分

        Args:
            previous_result: 前一次导出结果
            updated_document: 更新后的文档
            layout_engine: 排版引擎
            changes: 变更描述列表

        Returns:
            ExportResult: 新版本的导出结果
        """
        new_version = previous_result.version + 1
        return self.export_all(
            document=updated_document,
            layout_engine=layout_engine,
            version=new_version,
            change_log=changes,
        )

    def _export_pdf(self, html_content: str, output_dir: Path) -> Optional[Path]:
        """HTML → PDF 转换"""
        try:
            from weasyprint import HTML
            pdf_path = output_dir / "final_plan.pdf"
            HTML(string=html_content).write_pdf(str(pdf_path))
            return pdf_path
        except ImportError:
            raise
        except Exception as e:
            raise

    def _export_docx(self, markdown_text: str, project_name: str, output_dir: Path) -> Optional[Path]:
        """Markdown → Word (.docx) 转换"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '微软雅黑'
            font.size = Pt(12)

            lines = markdown_text.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()

                if not line:
                    i += 1
                    continue

                if line.startswith('# ') and not line.startswith('## '):
                    # 主标题
                    p = doc.add_heading(line[2:], level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    # 章标题
                    doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    # 节标题
                    doc.add_heading(line[4:], level=2)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('> '):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
                elif '|' in line and line.startswith('|'):
                    # 简单表格处理 - 获取跳过的索引
                    i = self._add_table_to_docx(doc, lines, i)
                    continue
                elif line == '---':
                    doc.add_paragraph('_' * 50)
                else:
                    doc.add_paragraph(line)

                i += 1

            docx_path = output_dir / "final_plan.docx"
            doc.save(str(docx_path))
            return docx_path

        except ImportError:
            raise
        except Exception as e:
            raise

    def _add_table_to_docx(self, doc, lines: List[str], start_idx: int) -> int:
        """简单表格添加到Word文档"""
        # 提取表格行
        table_lines = []
        for j in range(start_idx, len(lines)):
            line = lines[j].strip()
            if '|' in line:
                table_lines.append(line)
            else:
                break

        if len(table_lines) < 2:
            return start_idx + 1

        # 解析
        rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').replace(' ', '') == '' for c in cells):
                rows.append(cells)

        if not rows:
            return start_idx + len(table_lines)

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < len(table.rows[r_idx].cells):
                    table.rows[r_idx].cells[c_idx].text = cell_text

        doc.add_paragraph()  # 表后空行
        return start_idx + len(table_lines)

    def _generate_comprehensive_questionnaire(self, project_name: str, missing_sections: list) -> str:
        """生成通用客户补充资料引导问卷（覆盖所有缺失类型）"""
        # 按类别分组缺失项
        categories = {
            "技术类": [],
            "市场类": [],
            "产品类": [],
            "财务类": [],
            "规划类": [],
            "其他": [],
        }
        for m in missing_sections:
            if any(kw in m for kw in ["技术", "创新", "原理", "参数", "专利", "论文"]):
                categories["技术类"].append(m)
            elif any(kw in m for kw in ["市场", "客户", "竞品", "行业", "规模"]):
                categories["市场类"].append(m)
            elif any(kw in m for kw in ["产品", "样机", "规格", "型号", "场景"]):
                categories["产品类"].append(m)
            elif any(kw in m for kw in ["财务", "成本", "营收", "融资", "估值", "盈利"]):
                categories["财务类"].append(m)
            elif any(kw in m for kw in ["规划", "路线", "目标", "战略", "社会", "政策"]):
                categories["规划类"].append(m)
            else:
                categories["其他"].append(m)

        lines = [
            f"# [INFO] 客户补充资料引导问卷",
            f"",
            f"## 项目：{project_name}",
            f"",
            f"> 为将策划书从省赛金奖级别提升至国奖水准，请补充以下资料。",
            f"> 标注 🔴 的为关键项，补充后可显著提升评审得分。",
            f"",
            f"共 {len(missing_sections)} 项待补充，按类别整理如下：",
            f"",
        ]

        question_templates = {
            "核心技术原理描述": "🔴 请详细描述项目的核心技术原理（建议500字以上，含理论模型、工艺流程、关键参数）",
            "技术创新点列表": "🔴 请列出3-5个核心创新点，每个创新点附1-2段详细说明",
            "关键技术参数": "🔴 请补充关键技术性能参数表（参数名称 / 数值 / 测试条件 / 对标竞品）",
            "专利布局情况": "请列出已申请/授权的专利清单（名称/类型/申请号/状态）",
            "代表性论文": "请列出3-5篇代表性论文（标题/期刊/发表时间/作者）",
            "行业分析资料": "🔴 请提供行业分析资料（产业链位置、技术发展趋势、政策环境）",
            "市场调研数据": "🔴 请补充详细市场数据（TAM/SAM/SOM、客户画像、竞品分析）",
            "竞品对比数据": "请补充竞品对比的量化数据（关键参数、价格、市场份额）",
            "产品规格参数": "请补充产品详细规格表（型号/参数/工作条件/尺寸/重量）",
            "应用场景描述": "请描述3个典型应用场景（场景背景/客户痛点/解决方案/效果数据）",
            "产品迭代路线图": "请提供产品从当前版本到未来2-3代的演进路线",
            "成本结构数据": "🔴 请提供成本构成（研发/材料/制造/人力/营销等各项占比）",
            "收入预测依据": "请提供未来3年收入预测的依据（市场渗透率/定价策略/销售渠道）",
            "融资计划详情": "🔴 请明确融资额度、出让股权比例、资金用途分配",
            "技术发展路线": "请列出未来3-5年技术发展的关键里程碑",
            "社会价值量化": "请量化项目的预期社会价值（就业/环保/产业升级等）",
        }

        for category, items in categories.items():
            if not items:
                continue
            lines.append(f"## {category}（{len(items)}项）")
            lines.append("")
            for item in items:
                clean = item.replace("【待补充：", "").replace("】", "").replace("【待补充:", "").strip()
                question = question_templates.get(clean, f"请补充以下信息：{clean}")
                lines.append(f"- {question}")
            lines.append("")

        lines.append("---")
        lines.append("*填写完毕后，将本文件连同补充资料发送回智能体，系统将自动重新生成完整版策划书。*")
        return "\n".join(lines)

    def _generate_financial_questionnaire(self, project_name: str) -> str:
        """生成财务预测补充问卷"""
        return f"""# [INFO] 财务预测与融资计划补充问卷

## 项目：{project_name}

> 尊敬的客户，为完善策划书财务章节，请填写以下信息。已填写的内容将直接用于生成国奖标准的财务预测章节。

### 一、成本结构（请填写预估数值）
| 成本项目 | 第1年（万元） | 第2年（万元） | 第3年（万元） |
| :--- | :--- | :--- | :--- |
| 研发投入 | ____ | ____ | ____ |
| 原材料/生产 | ____ | ____ | ____ |
| 人力成本 | ____ | ____ | ____ |
| 市场推广 | ____ | ____ | ____ |
| 设备折旧 | ____ | ____ | ____ |
| 其他 | ____ | ____ | ____ |

### 二、收入预测
1. 核心产品单价：____ 元/件
2. 预计第1年销售量：____ 件 → 营收 ____ 万元
3. 预计第2年销售量：____ 件 → 营收 ____ 万元
4. 预计第3年销售量：____ 件 → 营收 ____ 万元
5. 收入增长的主要驱动力：________________

### 三、融资需求
1. 本轮融资额度：____ 万元
2. 出让股权比例：____ %
3. 投前估值：____ 万元（或PS/PE倍数为____倍）
4. 资金用途分配：
   - 研发投入 ____ %
   - 产能建设 ____ %
   - 市场拓展 ____ %
   - 团队扩充 ____ %
   - 其他 ____ %

### 四、盈利预测
1. 毛利率：第1年 ____ %，第2年 ____ %，第3年 ____ %
2. 净利率：第1年 ____ %，第2年 ____ %，第3年 ____ %
3. 预计盈亏平衡时间：第 ____ 年
4. 投资回收期：____ 年

### 五、未来规划
1. 3-5年技术里程碑（请列举）：
   - 2027年：________________
   - 2028年：________________
   - 2029年：________________
2. 市场拓展目标：第3年覆盖 ____ 个城市/省份
3. 预计带动就业：____ 人
4. 预计减少碳排放：____ 吨/年（如有）

---
*填写完毕后请回复本问卷，系统将自动重生成完整策划书。*
"""

    def _generate_privacy_statement(self, project_name: str) -> str:
        """生成数据隐私声明"""
        return f"""# 数据使用与隐私承诺

## 项目：{project_name}
## 生成日期：{datetime.now().strftime('%Y-%m-%d')}

### 我们对您的数据做出以下郑重承诺：

1. **专用性**：您提交的所有资料（包括但不限于项目简介、技术文档、专利证书、
   团队信息、合作协议等）仅用于本次策划书的生成，不会用于任何其他用途。

2. **不训练模型**：我们绝不以任何形式使用您的数据训练或改进AI模型。
   您的项目数据不会被用于增强任何公开或私有的AI模型。

3. **限期删除**：在任务交付后的{DATA_RETENTION_DAYS}日内，我们将物理删除
   您提交的所有原始资料及中间处理数据。

4. **数据隔离**：每个客户的项目数据在独立的存储空间中处理，
   不同客户之间的数据完全隔离。

5. **自主删除**：您可以随时要求立即删除您的所有数据，我们将在48小时内完成。

6. **不共享**：我们不会将您的数据分享、出售或透露给任何第三方。

### 如果您对数据使用有任何疑问或特殊要求，请联系我们。

---
*由全自动竞赛策划智能体自动生成*
"""
